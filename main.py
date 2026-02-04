from fastapi import FastAPI, UploadFile, File, Request, Form
from fastapi.responses import JSONResponse, HTMLResponse
from fastapi.templating import Jinja2Templates
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
import base64
import io
import zipfile
import re
import os
import httpx
from typing import Optional, List
from pathlib import Path
import uuid
import shutil
from PIL import Image

app = FastAPI(title="MilliyTest DOCX Parser API (Stable)")

# CORS sozlamasi
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Production da ma'lum domain'larni ko'rsatish
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["*"],  # Barcha header'larni ko'rsatish
)

# Templates uchun
templates = Jinja2Templates(directory="templates")

# API URL environment variable dan olish yoki default qiymat
API_URL = os.getenv("API_URL", "https://reyting.ideal-study.uz/api/public/tests")
QUESTIONS_API_URL = os.getenv("QUESTIONS_API_URL", "https://reyting.ideal-study.uz/api/questions")

# Rasmlar uchun vaqtinchalik papka
UPLOAD_DIR = Path("uploads/images")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


@app.get("/", response_class=HTMLResponse)
async def read_root(request: Request):
    """Asosiy sahifa."""
    return templates.TemplateResponse("index.html", {"request": request, "api_url": API_URL})


@app.get("/api/tests")
async def get_tests():
    """Barcha testlarni API dan olish."""
    try:
        async with httpx.AsyncClient() as client:
            response = await client.get(API_URL, timeout=10.0)
            response.raise_for_status()
            data = response.json()
            
            if data.get("success") and data.get("data"):
                # Testlarni qaytarish (faqat id va name)
                tests = [{"id": test.get("id"), "name": test.get("name")} for test in data["data"]]
                return JSONResponse({
                    "success": True,
                    "tests": tests
                })
            else:
                return JSONResponse({"success": False, "error": "Ma'lumot topilmadi"}, status_code=404)
                
    except httpx.HTTPError as e:
        print(f"❌ API xatolik: {e}")
        return JSONResponse({"success": False, "error": f"API ga ulanishda xatolik: {str(e)}"}, status_code=500)
    except Exception as e:
        print(f"❌ Umumiy xatolik: {e}")
        return JSONResponse({"success": False, "error": str(e)}, status_code=500)


@app.get("/api/test-data/{test_id}")
async def get_test_data(test_id: int):
    """Tanlangan test ma'lumotlarini API dan olish."""
    try:
        async with httpx.AsyncClient() as client:
            response = await client.get(API_URL, timeout=10.0)
            response.raise_for_status()
            data = response.json()
            
            if data.get("success") and data.get("data"):
                # Tanlangan testni topish
                test = next((t for t in data["data"] if t.get("id") == test_id), None)
                
                if test:
                    # Fanlar va sinflarni qaytarish
                    return JSONResponse({
                        "success": True,
                        "subjects": test.get("subjects", []),
                        "grades": test.get("grades", [])
                    })
                else:
                    return JSONResponse({"success": False, "error": "Test topilmadi"}, status_code=404)
            else:
                return JSONResponse({"success": False, "error": "Ma'lumot topilmadi"}, status_code=404)
                
    except httpx.HTTPError as e:
        print(f"❌ API xatolik: {e}")
        return JSONResponse({"success": False, "error": f"API ga ulanishda xatolik: {str(e)}"}, status_code=500)
    except Exception as e:
        print(f"❌ Umumiy xatolik: {e}")
        return JSONResponse({"success": False, "error": str(e)}, status_code=500)

def safe_read_zip(zf, path):
    """Zip ichidan faylni xavfsiz o‘qish (mavjud bo‘lmasa None)."""
    try:
        return zf.read(path)
    except KeyError:
        return None


def extract_crop_info(run_element):
    """Rasm elementidan crop ma'lumotlarini olish."""
    crop_info = None
    try:
        # a:srcRect elementini topish (crop ma'lumotlari)
        # Bir nechta joylarda qidirish
        src_rect = None
        
        # 1. To'g'ridan-to'g'ri a:srcRect
        src_rect = run_element.find(
            ".//{http://schemas.openxmlformats.org/drawingml/2006/main}srcRect"
        )
        
        # 2. Agar topilmasa, a:blip ichida qidirish
        if src_rect is None:
            blip = run_element.find(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
            )
            if blip is not None:
                # a:blip parent elementida qidirish
                parent = blip.getparent()
                if parent is not None:
                    src_rect = parent.find(
                        ".//{http://schemas.openxmlformats.org/drawingml/2006/main}srcRect"
                    )
        
        # 3. Agar hali ham topilmasa, a:pic ichida qidirish
        if src_rect is None:
            pic = run_element.find(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/picture}pic"
            )
            if pic is not None:
                src_rect = pic.find(
                    ".//{http://schemas.openxmlformats.org/drawingml/2006/main}srcRect"
                )
        
        if src_rect is not None:
            # Crop koordinatalari (EMU - English Metric Units)
            # 0-100000 aralig'ida, bu foiz sifatida ishlaydi
            l = float(src_rect.get("l", "0")) / 100000.0  # left
            t = float(src_rect.get("t", "0")) / 100000.0  # top
            r = float(src_rect.get("r", "0")) / 100000.0  # right
            b = float(src_rect.get("b", "0")) / 100000.0  # bottom
            
            # Agar crop mavjud bo'lsa
            if l > 0 or t > 0 or r > 0 or b > 0:
                crop_info = {"left": l, "top": t, "right": r, "bottom": b}
                print(f"📐 Crop ma'lumotlari topildi: l={l:.2%}, t={t:.2%}, r={r:.2%}, b={b:.2%}")
    except Exception as e:
        print(f"⚠️ Crop ma'lumotlarini o'qishda xatolik: {e}")
        import traceback
        traceback.print_exc()
    
    return crop_info


def crop_image(img_bytes, crop_info):
    """Rasmni crop qilish."""
    if not crop_info:
        return img_bytes
    
    try:
        # PIL Image ga o'tkazish
        img = Image.open(io.BytesIO(img_bytes))
        original_width, original_height = img.size
        
        # Crop koordinatalarini hisoblash
        left = int(original_width * crop_info["left"])
        top = int(original_height * crop_info["top"])
        right = int(original_width * (1 - crop_info["right"]))
        bottom = int(original_height * (1 - crop_info["bottom"]))
        
        # Koordinatalarni tekshirish
        left = max(0, min(left, original_width))
        top = max(0, min(top, original_height))
        right = max(left, min(right, original_width))
        bottom = max(top, min(bottom, original_height))
        
        # Crop qilish
        cropped_img = img.crop((left, top, right, bottom))
        
        # Bytes ga o'tkazish
        output = io.BytesIO()
        # Format ni saqlash
        img_format = img.format or "PNG"
        cropped_img.save(output, format=img_format)
        cropped_bytes = output.getvalue()
        
        print(f"✂️ Rasm crop qilindi: {original_width}x{original_height} -> {right-left}x{bottom-top}")
        
        return cropped_bytes
        
    except Exception as e:
        print(f"⚠️ Rasmni crop qilishda xatolik: {e}, original rasm qaytarilmoqda")
        return img_bytes


def extract_image_from_cell(cell, zip_content, image_index):
    """Cell ichidagi birinchi rasmni base64 qilib olish. a:blip yoki v:imagedata orqali."""
    try:
        found_image_ref = None
        run_element = None
        
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:

                # 1) a:blip orqali rasm izlash
                blips = run._element.findall(
                    ".//a:blip",
                    {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"},
                )
                for blip in blips:
                    embed_id = blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )
                    if embed_id:
                        found_image_ref = embed_id
                        run_element = run._element
                        break

                # 2) Agar blip topilmasa → v:imagedata orqali VML rasm qidirish
                if not found_image_ref:
                    v_imgs = run._element.findall(
                        ".//v:imagedata",
                        {"v": "urn:schemas-microsoft-com:vml"}
                    )
                    for vimg in v_imgs:
                        embed_id = vimg.get(
                            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
                        )
                        if embed_id:
                            found_image_ref = embed_id
                            run_element = run._element
                            break

                if found_image_ref:
                    break
            if found_image_ref:
                break

        if not found_image_ref:
            return None

        # 3) Crop ma'lumotlari
        crop_info = extract_crop_info(run_element) if run_element else None

        # 4) DOCX ZIPdan rasmni chiqarish
        with zipfile.ZipFile(io.BytesIO(zip_content), "r") as zf:
            rels_xml = safe_read_zip(zf, "word/_rels/document.xml.rels")
            if not rels_xml:
                return None
            rels_xml = rels_xml.decode("utf-8", errors="ignore")

            match = re.search(f'Id="{found_image_ref}".*?Target="([^"]*)"', rels_xml)
            if not match:
                return None

            media_file = f"word/{match.group(1)}"
            img_bytes = safe_read_zip(zf, media_file)
            if not img_bytes:
                return None

            # 5) WMF/EMF → PNG konvertatsiya (Word clipartlar uchun)
            ext = os.path.splitext(media_file)[1].lower()
            if ext in [".wmf", ".emf"]:
                try:
                    from PIL import Image
                    img = Image.open(io.BytesIO(img_bytes))
                    buffer = io.BytesIO()
                    img.save(buffer, format="PNG")
                    img_bytes = buffer.getvalue()
                except:
                    print("⚠ WMF/EMF konvert qilib bo'lmadi, o‘tkazib yuborildi.")
                    return None

            # 6) Crop qo‘llash
            if crop_info:
                img_bytes = crop_image(img_bytes, crop_info)

            # 7) MIME aniqlash
            mime_type = "image/png"
            if img_bytes[:4] == b'\x89PNG':
                mime_type = "image/png"
            elif img_bytes[:3] == b'\xFF\xD8\xFF':
                mime_type = "image/jpeg"
            else:
                if ext in [".jpg", ".jpeg"]:
                    mime_type = "image/jpeg"
                elif ext == ".gif":
                    mime_type = "image/gif"

            # 8) Base64
            img_base64 = base64.b64encode(img_bytes).decode()

            return f"data:{mime_type};base64,{img_base64}"

    except Exception as e:
        print(f"❗ Rasmni o'qishda xatolik: {e}")
        return None

def build_cell_data(cell, zip_content, image_index):
    """Cell uchun text/image obyektini tuzadi. 
    Agar text yo'q bo'lsa bo'sh string, agar image yo'q bo'lsa None qaytaradi."""
    text = ""
    try:
        text_val = cell.text.strip()
        text = text_val if text_val else ""
    except Exception:
        text = ""

    image_base64 = extract_image_from_cell(cell, zip_content, image_index)

    # Agar text ham image yo'q bo'lsa, None qaytarish (butun qatorni tashlash uchun)
    if not text and not image_base64:
        return None

    # Text bo'lsa image null, image bo'lsa text bo'sh string
    return {
        "text": text if text else "",
        "image": image_base64 if image_base64 else None
    }


async def _parse_and_send_one_file(
    content: bytes,
    test: Optional[str],
    language: Optional[str],
    class_id: Optional[str],
    subject: Optional[str],
) -> tuple:
    """Bitta DOCX kontentini parse qilib API ga yuboradi. Qaytaradi: (success, count, error_msg)."""
    image_files = []
    try:
        doc = Document(io.BytesIO(content))
        questions = []
        image_index = 0

        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) < 5:
                    continue

                # Har bir cell uchun data (har bir rasm uchun unique index)
                question_data = build_cell_data(row.cells[0], content, image_index)
                if question_data and question_data.get("image"):
                    image_index += 1

                correct_data = build_cell_data(row.cells[1], content, image_index)
                if correct_data and correct_data.get("image"):
                    image_index += 1

                wrong1_data = build_cell_data(row.cells[2], content, image_index)
                if wrong1_data and wrong1_data.get("image"):
                    image_index += 1

                wrong2_data = build_cell_data(row.cells[3], content, image_index)
                if wrong2_data and wrong2_data.get("image"):
                    image_index += 1

                wrong3_data = build_cell_data(row.cells[4], content, image_index)
                if wrong3_data and wrong3_data.get("image"):
                    image_index += 1

                # Savol yo'q bo'lsa — butun qatorni tashlaymiz
                if question_data is None:
                    continue

                # Javoblarni to'g'ri format qilish (None bo'lsa bo'sh object)
                def format_answer(answer_data):
                    """Javobni format qilish - None bo'lsa bo'sh object qaytaradi."""
                    if answer_data is None:
                        return {"text": "", "image": None}
                    return {
                        "text": answer_data.get("text", "") or "",
                        "image": answer_data.get("image") or None
                    }

                # Formatlangan savol (rasm base64 formatida)
                formatted_question = {
                    "question": {
                        "text": question_data.get("text", "") or "",
                        "image": question_data.get("image") or None
                    },
                    "correct": format_answer(correct_data),
                    "wrong1": format_answer(wrong1_data),
                    "wrong2": format_answer(wrong2_data),
                    "wrong3": format_answer(wrong3_data)
                }
                
                # Javoblarni tekshirish - agar bir xil javoblar bo'lsa, ularni takrorlamaslik
                correct_answer = formatted_question["correct"]
                wrong_answers = [
                    ("wrong1", formatted_question["wrong1"]),
                    ("wrong2", formatted_question["wrong2"]),
                    ("wrong3", formatted_question["wrong3"])
                ]
                
                # Correct javobni key sifatida saqlash
                correct_key = f"{correct_answer['text']}|{correct_answer['image']}"
                seen_answers = [correct_key]
                
                # Wrong javoblarni tekshirish
                for wrong_key, wrong_answer in wrong_answers:
                    answer_key = f"{wrong_answer['text']}|{wrong_answer['image']}"
                    
                    # Agar bu javob allaqachon ko'rilgan bo'lsa (correct yoki boshqa wrong bilan)
                    if answer_key in seen_answers:
                        # Agar bo'sh bo'lmasa, uni bo'sh qilish
                        if wrong_answer['text'] or wrong_answer['image']:
                            print(f"⚠️ Bir xil javob topildi ({wrong_key}): {answer_key[:50]}... - bo'sh qilindi")
                            formatted_question[wrong_key] = {"text": "", "image": None}
                        else:
                            # Bo'sh bo'lsa ham, key'ni qo'shish (takrorlanishni oldini olish uchun)
                            seen_answers.append(answer_key)
                    else:
                        # Yangi javob, key'ni qo'shish
                        seen_answers.append(answer_key)
                
                questions.append(formatted_question)

        # JSON formatda Laravel API ga yuborish (base64 rasmlar bilan)
        try:
            # Vaqtinchalik fayllarni o'chirish (endi kerak emas)
            for img_path in image_files:
                try:
                    if os.path.exists(img_path):
                        os.remove(img_path)
                except Exception as e:
                    print(f"⚠️ Faylni o'chirishda xatolik: {e}")

            async with httpx.AsyncClient(
                timeout=120.0,  # Base64 katta bo'lishi mumkin
                follow_redirects=True,  # 302 redirect'larni avtomatik kuzatish
                verify=False  # SSL sertifikat tekshiruvini o'chirish (kerak bo'lsa)
            ) as client:
                # JSON payload tayyorlash (Laravel questions ni JSON string kutadi)
                import json
                payload = {
                    "test_id": int(test) if test else None,
                    "language": language,
                    "grade_id": int(class_id) if class_id else None,
                    "subject_id": int(subject) if subject else None,
                    "questions": json.dumps(questions),
                }

                # Base64 rasmlar hajmini tekshirish
                total_size = 0
                for q in questions:
                    for key in ['question', 'correct', 'wrong1', 'wrong2', 'wrong3']:
                        if q.get(key, {}).get('image'):
                            total_size += len(q[key]['image'])
                
                print(f"📤 {len(questions)} ta savol yuborilmoqda...")
                print(f"📤 Base64 rasmlar umumiy hajmi: {total_size / 1024 / 1024:.2f} MB")

                # POST request (JSON format)
                response = await client.post(
                    QUESTIONS_API_URL,
                    json=payload,
                    headers={
                        "Content-Type": "application/json",
                        "Accept": "application/json",
                        "User-Agent": "FastAPI-DOCX-Parser/1.0"
                    }
                )
                
                # Response statusni tekshirish
                print(f"📡 Response status: {response.status_code}")
                print(f"📡 Response headers: {dict(response.headers)}")
                
                # 302 yoki 3xx status kod bo'lsa
                if response.status_code in [301, 302, 303, 307, 308]:
                    print(f"⚠️ Redirect detected: {response.headers.get('Location', 'N/A')}")
                    return (False, len(questions), f"Server redirect (Status: {response.status_code})")
                
                response.raise_for_status()
                api_response = response.json()

                # Vaqtinchalik fayllarni o'chirish
                for img_path in image_files:
                    try:
                        if os.path.exists(img_path):
                            os.remove(img_path)
                    except Exception as e:
                        print(f"⚠️ Faylni o'chirishda xatolik: {e}")

                return (True, len(questions), None)
        except httpx.HTTPStatusError as e:
            # HTTP xatolik (4xx, 5xx)
            # Xatolik bo'lsa ham fayllarni o'chirish
            for img_path in image_files:
                try:
                    if os.path.exists(img_path):
                        os.remove(img_path)
                except:
                    pass
            
            error_detail = f"Status: {e.response.status_code}"
            try:
                error_body = e.response.json()
                error_detail += f", Response: {error_body}"
            except:
                error_detail += f", Response: {e.response.text[:500]}"
            
            print(f"❌ API HTTP xatolik: {error_detail}")
            return (False, len(questions), f"API xatolik (Status: {e.response.status_code}): {str(e)}")
        except httpx.RequestError as e:
            # Network xatolik
            for img_path in image_files:
                try:
                    if os.path.exists(img_path):
                        os.remove(img_path)
                except:
                    pass
            
            print(f"❌ API ga ulanishda xatolik: {e}")
            return (False, 0, f"API ga ulanishda xatolik: {str(e)}")
        except Exception as e:
            # Boshqa xatoliklar
            for img_path in image_files:
                try:
                    if os.path.exists(img_path):
                        os.remove(img_path)
                except:
                    pass
            
            print(f"❌ API ga yuborishda noma'lum xatolik: {e}")
            import traceback
            traceback.print_exc()
            return (False, len(questions) if 'questions' in dir() else 0, str(e))

    except Exception as e:
        for img_path in image_files:
            try:
                if os.path.exists(img_path):
                    os.remove(img_path)
            except:
                pass
        print(f"❌ Server xatolik: {e}")
        import traceback
        traceback.print_exc()
        return (False, 0, str(e))


@app.post("/parse-docx/")
async def parse_docx(
    files: List[UploadFile] = File(...),
    test: Optional[str] = Form(None),
    language: Optional[str] = Form(None),
    class_id: Optional[str] = Form(None),
    subject: Optional[str] = Form(None),
):
    """Bir yoki bir nechta DOCX faylni ketma-ket parse qilib, har birini API ga yuboradi."""
    if not files:
        return JSONResponse(
            {"success": False, "error": "Kamida bitta fayl tanlang"},
            status_code=422,
        )
    total_questions = 0
    files_processed = 0
    files_failed = 0
    errors = []

    for idx, file in enumerate(files):
        try:
            content = await file.read()
            success, count, error_msg = await _parse_and_send_one_file(
                content, test, language, class_id, subject
            )
            if success:
                total_questions += count
                files_processed += 1
                print(f"✅ Fayl {idx + 1}/{len(files)}: {count} ta savol yuborildi")
            else:
                files_failed += 1
                err = error_msg or "Noma'lum xatolik"
                errors.append({"file": file.filename, "error": err})
                print(f"❌ Fayl {idx + 1}/{len(files)} ({file.filename}): {err}")
        except Exception as e:
            files_failed += 1
            errors.append({"file": file.filename, "error": str(e)})
            print(f"❌ Fayl {idx + 1}/{len(files)} ({file.filename}): {e}")

    return JSONResponse({
        "success": files_failed == 0,
        "total_questions": total_questions,
        "files_processed": files_processed,
        "files_failed": files_failed,
        "files_total": len(files),
        "errors": errors if errors else None,
        "message": f"{files_processed} ta fayl qayta ishlandi, {total_questions} ta savol yuborildi."
        + (f" {files_failed} ta faylda xatolik." if files_failed else ""),
    })